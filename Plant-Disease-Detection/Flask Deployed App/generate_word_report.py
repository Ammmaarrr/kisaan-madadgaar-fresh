from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    document = Document()

    # Title
    title = document.add_heading('Progress Report III: Integration of Advanced ML/DL or RL Model with Interpretability and Optimization', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    p = document.add_paragraph()
    p.add_run('Date: ').bold = True
    p.add_run('November 23, 2025\n')
    p.add_run('Project: ').bold = True
    p.add_run('Plant Disease Detection System')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Abstract
    document.add_heading('Abstract', level=1)
    document.add_paragraph(
        'This report details the integration of an advanced Convolutional Neural Network (CNN) for plant disease detection, '
        'focusing on model interpretability and optimization. We implemented a deep learning architecture to classify plant '
        'leaf diseases and integrated Grad-CAM (Gradient-weighted Class Activation Mapping) to provide visual explanations '
        'for the model\'s predictions. This ensures transparency and trust in the automated diagnosis system.'
    )

    # 1. Introduction
    document.add_heading('1. Introduction', level=1)
    document.add_paragraph(
        'Modern agriculture relies heavily on timely disease detection. Deep Learning (DL) models offer high accuracy but '
        'often lack interpretability, acting as "black boxes." This phase of the project addresses this limitation by '
        'integrating interpretability mechanisms into the deployed Flask application.'
    )

    # 2. Methodology
    document.add_heading('2. Methodology', level=1)

    document.add_heading('2.1 Model Architecture', level=2)
    document.add_paragraph(
        'We utilized a custom CNN architecture optimized for image classification tasks. The model consists of four '
        'convolutional blocks, each followed by ReLU activation, Batch Normalization, and Max Pooling. The dense layers '
        'include Dropout for regularization to prevent overfitting.'
    )
    
    p = document.add_paragraph()
    p.add_run('Key Components:').bold = True
    document.add_paragraph('Convolutional Layers: Extract spatial features from leaf images.', style='List Bullet')
    document.add_paragraph('Batch Normalization: Stabilizes learning and accelerates convergence.', style='List Bullet')
    document.add_paragraph('Dropout (0.4): Reduces overfitting by randomly deactivating neurons during training.', style='List Bullet')
    document.add_paragraph('Output Layer: 34 classes representing various plant diseases (Rice, Cotton, Wheat, Mango, Tomato, Potato, Pepper) and healthy states - optimized for Pakistan agricultural context.', style='List Bullet')

    document.add_heading('2.2 Optimization', level=2)
    document.add_paragraph('The model training process was optimized using:')
    document.add_paragraph('Loss Function: Cross-Entropy Loss for multi-class classification.', style='List Bullet')
    document.add_paragraph('Optimizer: Adam optimizer, chosen for its adaptive learning rate capabilities, ensuring efficient convergence.', style='List Bullet')
    document.add_paragraph('Data Augmentation: Applied during training to enhance model robustness against variations in lighting and orientation.', style='List Bullet')

    document.add_heading('2.3 Interpretability (Grad-CAM)', level=2)
    document.add_paragraph(
        'To provide insights into the model\'s decision-making process, we implemented Grad-CAM. This technique utilizes '
        'the gradients of the target concept flowing into the final convolutional layer to produce a coarse localization '
        'map highlighting the important regions in the image for predicting the concept.'
    )

    # 3. Implementation Results
    document.add_heading('3. Implementation Results', level=1)
    document.add_paragraph(
        'The system successfully processes input images, predicts the disease class, and generates a heatmap overlay '
        'indicating the affected areas on the leaf.'
    )

    document.add_heading('3.1 Input Image', level=2)
    document.add_paragraph('The system accepts leaf images captured via mobile devices or uploaded through the web interface.')
    
    try:
        document.add_picture('test_leaf.jpg', width=Inches(3.0))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = document.add_paragraph('Figure 1: Sample input leaf image.')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.style = 'Caption'
    except Exception as e:
        document.add_paragraph(f"[Image missing: {e}]")

    document.add_heading('3.2 Model Prediction & Interpretability', level=2)
    document.add_paragraph(
        'The model analyzes the image and the Grad-CAM algorithm generates a heatmap. The red regions in the heatmap '
        'correspond to the areas that contributed most significantly to the model\'s prediction, effectively highlighting '
        'the disease symptoms.'
    )

    try:
        document.add_picture('gradcam_output.png', width=Inches(3.0))
        last_paragraph = document.paragraphs[-1] 
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = document.add_paragraph('Figure 2: Grad-CAM visualization overlay. Red areas indicate high importance for the model\'s classification.')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.style = 'Caption'
    except Exception as e:
        document.add_paragraph(f"[Image missing: {e}]")

    # 4. Conclusion
    document.add_heading('4. Conclusion', level=1)
    document.add_paragraph(
        'The integration of the CNN model with Grad-CAM interpretability enhances the diagnostic capability of the Plant '
        'Disease Detection system. Users can now not only receive a diagnosis but also visually verify the symptoms '
        'detected by the AI, increasing confidence in the system\'s reliability.'
    )

    # 5. Future Work
    document.add_heading('5. Future Work', level=1)
    document.add_paragraph('Further optimization of model size for mobile deployment.', style='List Bullet')
    document.add_paragraph('Integration of Reinforcement Learning (RL) for adaptive treatment recommendations.', style='List Bullet')
    document.add_paragraph('Expansion of the dataset to cover more regional plant varieties.', style='List Bullet')

    document.save('Progress_Report_III.docx')
    print("Word document generated successfully.")

if __name__ == "__main__":
    create_report()
